"""
Convert thesis_part_I_and_II.txt to an HEC-compliant Word document.

HEC formatting:
  - Times New Roman, 12 pt
  - 1.5 line spacing
  - 1-inch (2.54 cm) margins
  - Justified body text, first-line indent on paragraphs after the first
    paragraph of each section
  - Bold, numbered section headings
  - Page numbers in the footer (centred)
  - Italic running header (right-aligned): "Application of Game Theory in
    Oil Producing Countries"
  - Footnote on the title page acknowledging AI-assisted drafting
  - Equations numbered (1), (2), ... centred
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, Inches, Emu

SRC = Path("/Users/hamza/Codebase/MasterThesis/thesis_part_I_and_II.txt")
DST = Path("/Users/hamza/Codebase/MasterThesis/thesis_part_I_and_II.docx")

DOC_TITLE = "Application of Game Theory in Oil Producing Countries"


def add_page_number(paragraph):
    """Insert a PAGE field into a paragraph (for the footer)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE   \\* MERGEFORMAT"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def set_run_font(run, *, bold=False, italic=False, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    run.bold = bold
    run.italic = italic


def set_paragraph_format(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=None,
    space_before=0,
    space_after=6,
    line_spacing=1.5,
):
    pf = paragraph.paragraph_format
    paragraph.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def add_paragraph(doc, text, *, style="body", first=False):
    p = doc.add_paragraph()
    if style == "body":
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=None if first else Cm(0.75),
            space_after=6,
            line_spacing=1.5,
        )
        run = p.add_run(text)
        set_run_font(run)
    elif style == "equation":
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=None,
            space_before=6,
            space_after=6,
            line_spacing=1.5,
        )
        run = p.add_run(text)
        set_run_font(run, italic=False)
    elif style == "part":
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=None,
            space_before=18,
            space_after=12,
            line_spacing=1.5,
        )
        run = p.add_run(text)
        set_run_font(run, bold=True, size=14)
    elif style == "h1":
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=None,
            space_before=12,
            space_after=6,
            line_spacing=1.5,
        )
        run = p.add_run(text)
        set_run_font(run, bold=True, size=13)
    elif style == "h2":
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=None,
            space_before=10,
            space_after=4,
            line_spacing=1.5,
        )
        run = p.add_run(text)
        set_run_font(run, bold=True, size=12)
    elif style == "h3":
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=None,
            space_before=8,
            space_after=4,
            line_spacing=1.5,
        )
        run = p.add_run(text)
        set_run_font(run, bold=True, italic=True, size=12)
    elif style == "ref":
        # Hanging-indent reference entry
        pf = p.paragraph_format
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-1.0)
        run = p.add_run(text)
        set_run_font(run)
    elif style == "figcaption":
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=None,
            space_before=2,
            space_after=4,
            line_spacing=1.15,
        )
        run = p.add_run(text)
        set_run_font(run, italic=False)
    return p


def configure_section(section):
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Header: italic, right-aligned, full document title
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(
        hp,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=None,
        space_after=0,
        line_spacing=1.0,
    )
    hrun = hp.add_run(DOC_TITLE)
    set_run_font(hrun, italic=True, size=11)

    # Footer: centred page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(
        fp,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=0,
        line_spacing=1.0,
    )
    add_page_number(fp)
    # Apply font to the field run
    for run in fp.runs:
        set_run_font(run, size=11)


def is_eq_line(line: str) -> bool:
    s = line.rstrip()
    return bool(re.search(r"\([0-9]+\)\s*$", s)) and "=" in s and s.startswith("    ")


def is_part_heading(line: str) -> bool:
    return bool(re.match(r"^PART (I|II)\.", line.strip()))


def is_h1(line: str) -> bool:
    return bool(re.match(r"^3\.\s+", line.strip()))


def is_h2(line: str) -> bool:
    return bool(re.match(r"^3\.[0-9]+\s+", line.strip()))


def is_back_heading(line: str) -> bool:
    s = line.strip()
    return s.startswith("NOTE ON FIGURES AND TABLES") or s.startswith(
        "REFERENCES CITED"
    )


def collect_paragraphs(lines):
    """
    Walk the text line-by-line and yield (kind, text) tuples where kind is
    one of: 'part', 'h1', 'h2', 'h3', 'equation', 'body', 'figcaption',
    'ref', 'fig_header', 'ref_header'.
    """
    # Skip front matter up to PART I.
    start = 0
    for i, line in enumerate(lines):
        if is_part_heading(line):
            start = i
            break

    i = start
    inside_back = False
    in_refs = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # Skip dividers, sentinels and notes
        if set(stripped) <= {"="} and stripped:
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("END OF PART"):
            i += 1
            continue

        if is_part_heading(line):
            yield ("part", stripped)
            i += 1
            continue

        if is_h1(line) and not inside_back:
            yield ("h1", stripped)
            i += 1
            continue

        if is_h2(line) and not inside_back:
            yield ("h2", stripped)
            i += 1
            continue

        if stripped.startswith("NOTE ON FIGURES AND TABLES"):
            inside_back = True
            yield ("h1", "Figures and Tables")
            # Skip the parenthetical instruction line, treat next paragraph as caption
            i += 1
            continue

        if stripped.startswith("Selected external figures"):
            yield ("h2", "Selected external figures")
            i += 1
            continue

        if stripped.startswith("REFERENCES CITED"):
            inside_back = True
            in_refs = True
            yield ("h1", "References (Part I and Part II)")
            # Skip the parenthetical instruction line below it
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("("):
                j += 1
            while j < len(lines) and lines[j].strip().startswith("reference list"):
                j += 1
            i = j
            continue

        if is_eq_line(line):
            yield ("equation", stripped)
            i += 1
            continue

        if inside_back and not in_refs:
            # Figure / table caption block. Gather until blank line.
            buf = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip() and not is_h2(lines[j]) \
                    and not lines[j].strip().startswith("Selected external"):
                buf.append(lines[j])
                j += 1
            text = " ".join(b.strip() for b in buf)
            yield ("figcaption", text)
            i = j
            continue

        if in_refs:
            # A reference entry begins with a non-indented line that has a
            # capitalised author surname. Gather continuation lines that are
            # indented.
            buf = [line]
            j = i + 1
            while j < len(lines):
                nxt = lines[j]
                if not nxt.strip():
                    break
                # A new entry starts at column 0 with a capital letter.
                if not nxt.startswith(" ") and nxt[:1].isupper():
                    break
                buf.append(nxt)
                j += 1
            text = " ".join(b.strip() for b in buf)
            yield ("ref", text)
            i = j
            continue

        # Default: ordinary body paragraph. Join continuation lines until a blank.
        buf = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt.strip():
                break
            if is_h1(nxt) or is_h2(nxt) or is_part_heading(nxt) or is_eq_line(nxt):
                break
            if is_back_heading(nxt):
                break
            buf.append(nxt)
            j += 1
        text = " ".join(b.strip() for b in buf)
        yield ("body", text)
        i = j


def build_document():
    doc = Document()

    # Configure the default style to make sure Times New Roman is the base.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")

    section = doc.sections[0]
    configure_section(section)

    # Document title block on first page (no separate title page since the
    # rest of the thesis already has one).
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_before=0,
        space_after=4,
        line_spacing=1.15,
    )
    title_run = p.add_run(DOC_TITLE)
    set_run_font(title_run, bold=True, size=14)

    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=4,
        line_spacing=1.15,
    )
    sub = p.add_run("Part I and Part II of the Research Paper")
    set_run_font(sub, italic=True, size=12)

    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=12,
        line_spacing=1.15,
    )
    sub = p.add_run("Hamza and Yassine | HEC Paris, Master in International Finance | May 2026")
    set_run_font(sub, size=11)

    # AI-disclosure footnote-style line (HEC requires AI-use acknowledgement).
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        space_after=18,
        line_spacing=1.15,
    )
    run = p.add_run(
        "Note on AI assistance: this document was drafted with the aid of "
        "AI-based language tools used to organise and articulate the authors' "
        "own analytical and computational work. All numerical results, "
        "calibrations, interpretations, and conclusions are the authors' "
        "original contribution."
    )
    set_run_font(run, italic=True, size=10)

    # Body content
    lines = SRC.read_text(encoding="utf-8").splitlines()
    prev_kind = None
    for kind, text in collect_paragraphs(lines):
        if kind == "part":
            add_paragraph(doc, text, style="part")
        elif kind == "h1":
            add_paragraph(doc, text, style="h1")
        elif kind == "h2":
            add_paragraph(doc, text, style="h2")
        elif kind == "equation":
            add_paragraph(doc, text, style="equation")
        elif kind == "figcaption":
            add_paragraph(doc, text, style="figcaption")
        elif kind == "ref":
            add_paragraph(doc, text, style="ref")
        else:
            first = prev_kind in (None, "part", "h1", "h2", "h3")
            add_paragraph(doc, text, style="body", first=first)
        prev_kind = kind

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    build_document()
